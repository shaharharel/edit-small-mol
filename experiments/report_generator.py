import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict
from datetime import datetime
from scripts.visualization.plot_training_progress import plot_training_progress_per_epoch
from cluster_analysis import perform_cluster_analysis
from edit_embedding_comparison import compare_edit_embeddings


def create_comparison_plot(results: Dict, metric: str, output_path: str):
    methods = list(results.keys())
    properties = list(results[methods[0]].keys())

    data = []
    for method in methods:
        for prop in properties:
            if prop in results[method]:
                value = results[method][prop]['metrics'].get(metric, np.nan)
                data.append({
                    'Method': method,
                    'Property': prop,
                    'Value': value
                })

    df = pd.DataFrame(data)

    fig, ax = plt.subplots(figsize=(12, 6))

    x = np.arange(len(properties))
    width = 0.8 / len(methods)

    for i, method in enumerate(methods):
        method_data = df[df['Method'] == method]
        values = [method_data[method_data['Property'] == p]['Value'].values[0] if len(method_data[method_data['Property'] == p]) > 0 else 0 for p in properties]
        ax.bar(x + i * width, values, width, label=method)

    ax.set_xlabel('Property', fontsize=12)
    ax.set_ylabel(metric.upper(), fontsize=12)
    ax.set_title(f'{metric.upper()} Comparison Across Methods', fontsize=14, fontweight='bold')
    ax.set_xticks(x + width * (len(methods) - 1) / 2)
    ax.set_xticklabels(properties, rotation=45, ha='right')
    ax.legend()
    ax.grid(axis='y', alpha=0.3)

    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()

    return output_path


def create_scatter_plot(y_true, y_pred, title: str, output_path: str):
    fig, ax = plt.subplots(figsize=(8, 8))

    ax.scatter(y_true, y_pred, alpha=0.5, s=20)

    min_val = min(y_true.min(), y_pred.min())
    max_val = max(y_true.max(), y_pred.max())
    ax.plot([min_val, max_val], [min_val, max_val], 'r--', lw=2, label='Perfect Prediction')

    ax.set_xlabel('True Values', fontsize=12)
    ax.set_ylabel('Predicted Values', fontsize=12)
    ax.set_title(title, fontsize=14, fontweight='bold')
    ax.legend()
    ax.grid(alpha=0.3)

    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()

    return output_path


def create_metrics_table(results: Dict, dataset_name: str) -> pd.DataFrame:
    rows = []
    for method_name, method_results in results.items():
        for prop, prop_results in method_results.items():
            metrics = prop_results['metrics']
            rows.append({
                'Method': method_name,
                'Property': prop,
                'MAE': f"{metrics.get('mae', np.nan):.4f}",
                'RMSE': f"{metrics.get('rmse', np.nan):.4f}",
                'R²': f"{metrics.get('r2', np.nan):.4f}",
                'Pearson': f"{metrics.get('pearson_r', np.nan):.4f}",
                'Spearman': f"{metrics.get('spearman_r', np.nan):.4f}"
            })

    return pd.DataFrame(rows)


def generate_report(results: Dict, config, trained_models: Dict = None, embeddings: Dict = None, train_data: Dict = None) -> str:
    output_dir = Path(config.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_name = f"{config.experiment_name}_{timestamp}"

    # Create separate directories for reports and images
    reports_dir = output_dir / "reports"
    images_dir = output_dir / "images" / timestamp
    reports_dir.mkdir(parents=True, exist_ok=True)
    images_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    title = doc.add_heading(f'Experiment Report: {config.experiment_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('Experiment Configuration', 1)

    # General Configuration
    doc.add_heading('General Settings', 2)
    doc.add_paragraph(f"Experiment Name: {config.experiment_name}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Data File: {config.data_file}")
    doc.add_paragraph(f"Random Seed: {config.random_seed}")

    # Data Configuration
    doc.add_heading('Data Configuration', 2)
    doc.add_paragraph(f"Number of Tasks: {config.num_tasks}")
    doc.add_paragraph(f"Minimum Pairs per Property: {config.min_pairs_per_property}")
    doc.add_paragraph(f"Splitter Type: {config.splitter_type}")
    doc.add_paragraph(f"Train/Val/Test Ratio: {config.train_ratio}/{config.val_ratio}/{config.test_ratio}")
    if config.splitter_params and config.splitter_type in config.splitter_params:
        params = config.splitter_params[config.splitter_type]
        doc.add_paragraph(f"Splitter Parameters: {params}")

    # Few-shot metadata (if applicable)
    if hasattr(config, 'few_shot_metadata') and config.few_shot_metadata:
        doc.add_heading('Few-Shot Learning Configuration', 3)
        metadata = config.few_shot_metadata
        doc.add_paragraph(f"Few-Shot Fraction: {metadata['few_shot_fraction']*100:.0f}%")
        doc.add_paragraph(f"Few-Shot Samples per Property: {metadata['few_shot_samples']}")
        doc.add_paragraph(f"\nFew-Shot Properties ({len(metadata['few_shot_properties'])}):")
        for prop in metadata['few_shot_properties']:
            doc.add_paragraph(f"  • {prop}", style='List Bullet')
        doc.add_paragraph(f"\nRegular Properties ({len(metadata['regular_properties'])}):")
        for prop in metadata['regular_properties']:
            doc.add_paragraph(f"  • {prop}", style='List Bullet')

    # Embedding Configuration
    doc.add_heading('Embedding Configuration', 2)
    doc.add_paragraph(f"Embedder Type: {config.embedder_type}")

    # Methods Configuration
    doc.add_heading('Methods Configuration', 2)
    for i, method in enumerate(config.methods, 1):
        doc.add_paragraph(f"{i}. {method['name']}", style='List Bullet')
        doc.add_paragraph(f"   Type: {method['type']}")
        doc.add_paragraph(f"   Hidden Dimensions: {method.get('hidden_dims', 'N/A')}")
        doc.add_paragraph(f"   Dropout: {method.get('dropout', 'N/A')}")
        doc.add_paragraph(f"   Learning Rate: {method.get('lr', 'N/A')}")
        doc.add_paragraph(f"   Batch Size: {method.get('batch_size', 'N/A')}")
        doc.add_paragraph(f"   Max Epochs: {method.get('max_epochs', 'N/A')}")
        if method['type'] == 'edit_framework':
            doc.add_paragraph(f"   Use Edit Fragments: {method.get('use_edit_fragments', False)}")
        if 'load_checkpoint' in method and method['load_checkpoint']:
            doc.add_paragraph(f"   Loaded from Checkpoint: {method['load_checkpoint']}")

    # Analysis Configuration
    doc.add_heading('Analysis Configuration', 2)
    doc.add_paragraph(f"Metrics: {', '.join(config.metrics)}")
    doc.add_paragraph(f"Include Cluster Analysis: {config.include_cluster_analysis}")
    if config.include_cluster_analysis:
        doc.add_paragraph(f"Number of Clusters: {config.n_clusters}")
    doc.add_paragraph(f"Include Edit Embedding Comparison: {config.include_edit_embedding_comparison}")
    doc.add_paragraph(f"Save Models: {config.save_models}")

    if trained_models:
        doc.add_heading('Training Progress', 1)

        task_names = results.get('test', {}).get(list(results.get('test', {}).keys())[0] if results.get('test') else '', {}).keys() if results.get('test') else []
        task_names = list(task_names) if task_names else []

        for method_name, method_info in trained_models.items():
            model = method_info['model']

            try:
                fig, n_epochs = plot_training_progress_per_epoch(
                    model=model,
                    task_names=task_names,
                    steps_per_epoch=None
                )

                if fig is not None:
                    doc.add_heading(f'{method_name}', 2)

                    plot_path = images_dir / f"training_{method_name}.png"
                    fig.savefig(str(plot_path), dpi=150, bbox_inches='tight')
                    plt.close(fig)

                    doc.add_picture(str(plot_path), width=Inches(6))
                    doc.add_paragraph(f"Training completed in {n_epochs} epochs")
            except Exception as e:
                doc.add_heading(f'{method_name}', 2)
                doc.add_paragraph(f"Training history not available (model was loaded from checkpoint)")

    doc.add_heading('Test Set Results', 1)

    test_results = results['test']

    for metric in config.metrics:
        plot_path = images_dir / f"test_{metric}_comparison.png"
        create_comparison_plot(test_results, metric, str(plot_path))

        doc.add_heading(f'{metric.upper()} Comparison', 2)
        doc.add_picture(str(plot_path), width=Inches(6))

    doc.add_heading('Detailed Metrics Tables', 2)

    all_properties = set()
    for method_results in test_results.values():
        all_properties.update(method_results.keys())

    for prop in sorted(all_properties):
        doc.add_heading(f'Property: {prop}', 3)

        prop_rows = []
        for method_name, method_results in test_results.items():
            if prop in method_results:
                metrics = method_results[prop]['metrics']
                prop_rows.append({
                    'Method': method_name,
                    'MAE': f"{metrics.get('mae', np.nan):.4f}",
                    'RMSE': f"{metrics.get('rmse', np.nan):.4f}",
                    'R²': f"{metrics.get('r2', np.nan):.4f}",
                    'Pearson': f"{metrics.get('pearson_r', np.nan):.4f}",
                    'Spearman': f"{metrics.get('spearman_r', np.nan):.4f}"
                })

        if prop_rows:
            prop_df = pd.DataFrame(prop_rows)

            table = doc.add_table(rows=len(prop_df) + 1, cols=len(prop_df.columns))
            table.style = 'Light Grid Accent 1'

            for i, col in enumerate(prop_df.columns):
                table.rows[0].cells[i].text = col
                table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

            for i, row in prop_df.iterrows():
                for j, value in enumerate(row):
                    table.rows[i + 1].cells[j].text = str(value)

    doc.add_heading('Prediction Scatter Plots', 1)

    for method_name, method_results in test_results.items():
        doc.add_heading(f'Method: {method_name}', 2)

        properties = list(method_results.keys())
        plot_paths = []

        for prop, prop_results in method_results.items():
            y_true = prop_results['y_true']
            y_pred = prop_results['y_pred']

            plot_path = images_dir / f"scatter_{method_name}_{prop}.png"
            create_scatter_plot(
                y_true, y_pred,
                f"{method_name} - {prop}",
                str(plot_path)
            )
            plot_paths.append((prop, str(plot_path)))

        plots_per_row = 4
        num_plots = len(plot_paths)
        num_rows = (num_plots + plots_per_row - 1) // plots_per_row

        grid_table = doc.add_table(rows=num_rows, cols=plots_per_row)
        grid_table.autofit = False

        for idx, (prop, plot_path) in enumerate(plot_paths):
            row_idx = idx // plots_per_row
            col_idx = idx % plots_per_row
            cell = grid_table.rows[row_idx].cells[col_idx]

            paragraph = cell.paragraphs[0]
            paragraph.text = prop
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0]
            run.font.size = Pt(8)
            run.font.bold = True

            cell.add_paragraph()
            run = cell.paragraphs[1].add_run()
            run.add_picture(plot_path, width=Inches(0.5))

    if results['additional_test']:
        doc.add_heading('Additional Test Set Results', 1)

        for test_name, test_results in results['additional_test'].items():
            doc.add_heading(f'Dataset: {test_name}', 2)

            if test_results:
                metrics_df = create_metrics_table(test_results, test_name)

                table = doc.add_table(rows=len(metrics_df) + 1, cols=len(metrics_df.columns))
                table.style = 'Light Grid Accent 1'

                for i, col in enumerate(metrics_df.columns):
                    table.rows[0].cells[i].text = col
                    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

                for i, row in metrics_df.iterrows():
                    for j, value in enumerate(row):
                        table.rows[i + 1].cells[j].text = str(value)

    if config.include_cluster_analysis and trained_models and embeddings and train_data:
        doc.add_page_break()
        doc.add_heading('Edit Embedding Cluster Analysis', 1)

        for method_name, method_info in trained_models.items():
            if method_info['type'] == 'edit_framework':
                try:
                    task_names = list(train_data.keys())
                    test_df_list = [train_data[prop]['test'] for prop in task_names]
                    test_combined = pd.concat(test_df_list, ignore_index=True)

                    mol_emb_a_test_list = [embeddings[prop]['test']['mol_a'] for prop in task_names]
                    mol_emb_b_test_list = [embeddings[prop]['test']['mol_b'] for prop in task_names]
                    mol_emb_a_test = np.vstack(mol_emb_a_test_list)
                    mol_emb_b_test = np.vstack(mol_emb_b_test_list)

                    fig, cluster_stats = perform_cluster_analysis(
                        model=method_info['model'],
                        test_df=test_combined,
                        mol_emb_a_test=mol_emb_a_test,
                        mol_emb_b_test=mol_emb_b_test,
                        n_clusters=config.n_clusters,
                        device=method_info['model'].device
                    )

                    doc.add_heading(f'{method_name} - K-Means Clustering', 2)

                    plot_path = images_dir / f"cluster_{method_name}.png"
                    fig.savefig(str(plot_path), dpi=150, bbox_inches='tight')
                    plt.close(fig)

                    doc.add_picture(str(plot_path), width=Inches(6))

                    doc.add_heading('Cluster Statistics', 3)
                    for cluster_id, stats in cluster_stats.items():
                        doc.add_paragraph(f"Cluster {cluster_id}: {stats['size']} edits ({stats['pct']:.1f}%)")

                except Exception as e:
                    doc.add_paragraph(f"Cluster analysis not available: {str(e)}")

    if config.include_edit_embedding_comparison and trained_models and embeddings and train_data:
        doc.add_page_break()
        doc.add_heading('Edit Embedding Comparison', 1)

        for method_name, method_info in trained_models.items():
            if method_info['type'] == 'edit_framework':
                try:
                    task_names = list(train_data.keys())
                    test_df_list = [train_data[prop]['test'] for prop in task_names]
                    test_combined = pd.concat(test_df_list, ignore_index=True)

                    mol_emb_a_test_list = [embeddings[prop]['test']['mol_a'] for prop in task_names]
                    mol_emb_b_test_list = [embeddings[prop]['test']['mol_b'] for prop in task_names]
                    mol_emb_a_test = np.vstack(mol_emb_a_test_list)
                    mol_emb_b_test = np.vstack(mol_emb_b_test_list)

                    figures, results_df = compare_edit_embeddings(
                        model=method_info['model'],
                        test_df=test_combined,
                        mol_emb_a_test=mol_emb_a_test,
                        mol_emb_b_test=mol_emb_b_test,
                        task_names=task_names,
                        device=method_info['model'].device
                    )

                    doc.add_heading(f'{method_name} - Baseline vs Trained Embeddings', 2)

                    plot_names = ['r2_comparison', 'r2_improvement', 'mae_improvement']
                    for i, fig in enumerate(figures):
                        plot_path = images_dir / f"emb_comp_{method_name}_{plot_names[i]}.png"
                        fig.savefig(str(plot_path), dpi=150, bbox_inches='tight')
                        plt.close(fig)

                        doc.add_picture(str(plot_path), width=Inches(6))

                except Exception as e:
                    doc.add_paragraph(f"Edit embedding comparison not available: {str(e)}")

    doc.add_page_break()
    doc.add_heading('Summary', 1)
    doc.add_paragraph(f"Experiment completed successfully on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Results saved to: {output_dir}")
    doc.add_paragraph(f"Images saved to: {images_dir}")

    report_path = reports_dir / f"{config.experiment_name}.docx"
    doc.save(str(report_path))

    print(f"\n{'='*80}")
    print(f"Report generated: {report_path}")
    print(f"Images saved to: {images_dir}")
    print(f"{'='*80}")

    return str(report_path)
